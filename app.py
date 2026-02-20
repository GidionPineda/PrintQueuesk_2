from flask import Flask, render_template, request, redirect, url_for, send_from_directory
from flask_socketio import SocketIO, emit
from flask_cors import CORS
import os
import fitz  # PyMuPDF for PDF processing
from docx import Document  # For handling Word documents
import qrcode
from config.firebase_config import db
from werkzeug.utils import secure_filename
from datetime import datetime
import pytz

# Set Philippines timezone
ph_timezone = pytz.timezone('Asia/Manila')
def get_local_time():
    """Get current time in Philippines timezone"""
    return datetime.now(ph_timezone).strftime('%Y-%m-%d %H:%M:%S')

# Flask application setup
app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024  # 100 MB limit
CORS(app)  # Enable CORS for all routes
socketio = SocketIO(app, cors_allowed_origins="*")

# Allowed file extensions
ALLOWED_EXTENSIONS = {'docx', 'pdf'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def get_total_pages(file_path):
    try:
        print(f"[DEBUG] get_total_pages called for: {file_path}")
        
        if file_path.lower().endswith('.pdf'):
            with fitz.open(file_path) as pdf:
                page_count = pdf.page_count
                print(f"[DEBUG] PDF page count: {page_count}")
                return page_count
                
        elif file_path.lower().endswith(('.docx')):
            try:
                # Open DOCX file in read-only mode to avoid locking
                with open(file_path, 'rb') as f:
                    doc = Document(f)
                # Count only manual page breaks
                page_breaks = 0
                for para in doc.paragraphs:
                    for run in para.runs:
                        xml = run._element.xml
                        if 'w:br' in xml and 'type="page"' in xml:
                            page_breaks += 1
                if page_breaks > 0:
                    page_count = page_breaks + 1
                    print(f"[DEBUG] DOCX manual page breaks found: {page_breaks}, page count: {page_count}")
                    return page_count
                # Fallback: Estimate by paragraphs only if no page breaks
                para_count = len(doc.paragraphs)
                page_estimate = max(1, (para_count // 33) + (1 if para_count % 33 else 0))
                print(f"[DEBUG] DOCX fallback paragraph count: {para_count}, estimated pages: {page_estimate}")
                return page_estimate
            except Exception as fallback_error:
                print(f"[DEBUG] python-docx failed: {fallback_error}")
                return 1
            
    except Exception as e:
        print(f"[DEBUG] Error in get_total_pages: {e}")
        return 1  # Default to 1 page instead of None

# Ensure 'uploads' directory exists
if not os.path.exists('uploads'):
    os.makedirs('uploads')

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/uploads/<filename>')
def uploaded_file(filename):
    uploads_dir = os.path.join(os.path.dirname(__file__), 'uploads')
    return send_from_directory(uploads_dir, filename)

@app.route('/upload', methods=['GET', 'POST'])
def upload_file():
    if request.method == 'GET':
        return render_template('index.html')
    if 'file' not in request.files:
        return redirect(request.url)
    file = request.files['file']
    if file.filename == '':
        return "No selected file.", 400
    
    # Check if file type is allowed
    if not file or not allowed_file(file.filename):
        return "INVALID FILE TYPE: Only PDF, and DOCX files are allowed.", 400
    
    # Check file size (100 MB limit)
    file.seek(0, os.SEEK_END)
    file_length = file.tell()
    file.seek(0)  # Reset file pointer
    max_size = 100 * 1024 * 1024  # 100 MB in bytes
    if file_length > max_size:
        return f"FILE TOO LARGE: Maximum file size is 100 MB. Your file is {file_length/1024/1024:.1f} MB.", 400
    
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        uploads_dir = os.path.join(os.path.dirname(__file__), 'uploads')
        if not os.path.exists(uploads_dir):
            os.makedirs(uploads_dir)
        local_path = os.path.join(uploads_dir, filename)
        file.save(local_path)
        file_size = os.path.getsize(local_path)
        total_pages = get_total_pages(local_path)
        if total_pages is None:
            os.remove(local_path)
            return "Error processing the file.", 500
        try:
            # Build download URL for the uploaded file
            download_url = url_for('uploaded_file', filename=filename, _external=True)
            # Store metadata in Firebase Realtime Database
            jobs_ref = db.reference('jobs/print_jobs')
            now_ph = get_local_time()
            job_data = {
                'file_name': filename,
                'file_size': file_size,
                'total_pages': total_pages,
                'status': 'pending',
                'created_at': now_ph,
                'local_path': local_path,
                'download_url': download_url,
                'details': [
                    {
                        'file_name': filename,
                        'page_range': 'all',
                        'color_mode': '',
                        'total_price': 0,
                        'inserted_amount': 0,
                        'change_amount': 0,
                        'status': 'pending',
                        'created_at': now_ph,
                        'local_path': local_path,
                        'download_url': download_url
                    }
                ]
            }
            # Auto-increment Job ID logic
            existing_jobs = jobs_ref.get()
            if existing_jobs:
                # Handle both dict and list cases
                if isinstance(existing_jobs, dict):
                    existing_ids = [int(k) for k in existing_jobs.keys() if str(k).isdigit()]
                elif isinstance(existing_jobs, list):
                    existing_ids = [i for i in range(len(existing_jobs))]
                else:
                    existing_ids = []
                
                next_job_id = str(max(existing_ids) + 1) if existing_ids else '1'
            else:
                next_job_id = '1'

            # Store job with numerical Job ID
            jobs_ref.child(next_job_id).set(job_data)

            # Emit SocketIO event for the GUI, now with download_url
            socketio.emit('file_uploaded', {
                'file_name': filename,
                'total_pages': total_pages,
                'job_id': next_job_id,
                'download_url': download_url
            })
            return render_template('uploaded_file.html', filename=filename, file_size=file_size, total_pages=total_pages)
        except Exception as e:
            print(f"Error: {e}")
            socketio.emit('file_status_update', {'document_name': filename, 'status': 'failed'})
            return f"Error processing file: {e}", 500
        # Do not delete the file here, since it needs to be downloadable
    return "Invalid file type. Please upload a .docx, or .pdf file."

@app.route('/generate_wifi_qr')
def generate_wifi_qr():
    ssid = "YourSSID"
    password = "YourPassword"

    wifi_config = f"WIFI:S:{ssid};T:WPA;P:{password};;"
    qr = qrcode.QRCode()
    qr.add_data(wifi_config)
    qr.make(fit=True)

    img = qr.make_image(fill="black", back_color="white")
    qr_code_path = os.path.join('static', 'wifi_qr_code.png')
    img.save(qr_code_path)
    
    return render_template('wifi_qr.html', qr_code_path=qr_code_path)

@app.route('/api/update_job_payment', methods=['POST'])
def update_job_payment():
    """API endpoint to receive payment data from the printing machine"""
    try:
        data = request.get_json()
        
        # Validate required fields
        if not data or 'job_id' not in data:
            return {"error": "Missing job_id"}, 400
        
        job_id = str(data['job_id'])
        inserted_amount = float(data.get('inserted_amount', 0))
        change_amount = float(data.get('change_amount', 0))
        total_price = float(data.get('total_price', 0))
        status = data.get('status', 'paid')
        
        print(f"[API] Received payment update for job_id: {job_id}")
        print(f"[API] inserted_amount: {inserted_amount}, change_amount: {change_amount}, total_price: {total_price}")
        
        # Update Firebase with payment information
        jobs_ref = db.reference('jobs/print_jobs')
        job_ref = jobs_ref.child(job_id)
        
        # Check if job exists
        job_data = job_ref.get()
        if not job_data:
            print(f"[API] Job {job_id} not found")
            return {"error": f"Job {job_id} not found"}, 404
        
        # Update the payment details
        now_ph = get_local_time()
        update_data = {
            'inserted_amount': inserted_amount,
            'change_amount': change_amount,
            'total_price': total_price,
            'status': status,
            'payment_updated_at': now_ph
        }
        
        # Update details array if it exists
        if 'details' in job_data and isinstance(job_data['details'], list) and len(job_data['details']) > 0:
            job_ref.child('details').child('0').update(update_data)
        
        # Also update root level status
        job_ref.update({
            'status': status,
            'updated_at': now_ph
        })
        
        print(f"[API] Successfully updated payment for job {job_id}")
        
        return {
            "success": True,
            "message": "Payment data updated successfully",
            "job_id": job_id,
            "data": update_data
        }, 200
        
    except Exception as e:
        print(f"[API] Error updating payment: {e}")
        import traceback
        traceback.print_exc()
        return {"error": str(e)}, 500

@app.teardown_appcontext
def close_db_connection(exception):
    pass  # No persistent connection to close with Firebase

@socketio.on('update_status')
def update_status(data):
    document_name = data['document_name']
    status = data['status']

    # Find the job by document_name and update status
    jobs_ref = db.reference('jobs/print_jobs')
    jobs = jobs_ref.order_by_child('file_name').equal_to(document_name).get()  # Changed from 'document_name' to 'file_name'
    
    if jobs:
        # Handle both dict and list cases
        if isinstance(jobs, dict):
            for job_id, job in jobs.items():
                jobs_ref.child(job_id).update({'status': status})
                # Also update the status in details if exists
                if 'details' in job:
                    if isinstance(job['details'], dict):
                        for detail_id, detail in job['details'].items():
                            jobs_ref.child(job_id).child('details').child(detail_id).update({'status': status})
                    elif isinstance(job['details'], list):
                        for i, detail in enumerate(job['details']):
                            if isinstance(detail, dict):
                                jobs_ref.child(job_id).child('details').child(str(i)).update({'status': status})
                socketio.emit('status_update', {'document_name': document_name, 'status': status})
                break
        elif isinstance(jobs, list):
            for job in jobs:
                if isinstance(job, dict) and job.get('file_name') == document_name:
                    job_id = job.get('id')  # You might need to adjust this based on your actual structure
                    if job_id:
                        jobs_ref.child(job_id).update({'status': status})
                        socketio.emit('status_update', {'document_name': document_name, 'status': status})
                        break

if __name__ == '__main__':
    # Get the port from the environment, defaulting to 5000 for local development
    port = int(os.environ.get('PORT', 5000))
    socketio.run(app, host='0.0.0.0', port=port, allow_unsafe_werkzeug=True)
