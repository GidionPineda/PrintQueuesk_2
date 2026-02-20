# hotspot_config.py
"""
Configuration and server logic for the hotspot file upload functionality.
This module handles the HTTP server and file upload processing.
"""

import os
import http.server
import socketserver
import threading
from datetime import datetime
import fitz  # PyMuPDF
from docx import Document
from firebase_admin import db
import time
import qrcode
from io import BytesIO
import pytz
import json


def get_total_pages(file_path):
    """Get total pages from PDF or DOCX files."""
    try:
        if file_path.lower().endswith('.pdf'):
            with fitz.open(file_path) as pdf:
                return pdf.page_count
                
        elif file_path.lower().endswith(('.doc', '.docx')):
            try:
                with open(file_path, 'rb') as f:
                    doc = Document(f)
                # Count manual page breaks
                page_breaks = 0
                for para in doc.paragraphs:
                    for run in para.runs:
                        xml = run._element.xml
                        if 'w:br' in xml and 'type="page"' in xml:
                            page_breaks += 1
                if page_breaks > 0:
                    return page_breaks + 1
                # Fallback: Estimate by paragraphs
                para_count = len(doc.paragraphs)
                return max(1, (para_count // 33) + (1 if para_count % 33 else 0))
            except Exception:
                return 1
    except Exception as e:
        print(f"Error getting total pages: {e}")
        return 1


def save_file_to_firebase(file_path, file_name):
    """Save file information to Firebase and return job data."""
    max_retries = 3
    retry_delay = 5  # seconds

    for attempt in range(max_retries):
        try:
            # On success, return the job_data
            job_data = _save_file_to_firebase_attempt(file_path, file_name)
            if job_data:
                return job_data
            # If _save_file_to_firebase_attempt returns None without an exception, it's a non-retryable error.
            else:
                print("Failed to save to Firebase due to an internal error without exception.")
                return None

        except Exception as e:
            print(f"Attempt {attempt + 1} of {max_retries} failed: {e}")
            if attempt < max_retries - 1:
                print(f"Retrying in {retry_delay} seconds...")
                time.sleep(retry_delay)
            else:
                print("All retries failed. Could not save to Firebase.")
                import traceback
                traceback.print_exc()
                return None


def _save_file_to_firebase_attempt(file_path, file_name):
    """Attempt to save file information to Firebase. Can raise exceptions."""
    total_pages = get_total_pages(file_path)
    file_size = os.path.getsize(file_path)
    
    # Set Philippines timezone
    ph_timezone = pytz.timezone('Asia/Manila')
    now_ph = datetime.now(ph_timezone).strftime("%Y-%m-%d %H:%M:%S")
    
    # Create a new job entry in Firebase with the same structure as the web app
    jobs_ref = db.reference('jobs/print_jobs')
    
    # Get existing jobs to auto-increment job ID
    existing_jobs = jobs_ref.get()
    if existing_jobs:
        if isinstance(existing_jobs, dict):
            existing_ids = [int(k) for k in existing_jobs.keys() if str(k).isdigit()]
        elif isinstance(existing_jobs, list):
            # Handle cases where Firebase returns a list for sequential keys
            existing_ids = [i for i, v in enumerate(existing_jobs) if v is not None]
        else:
            existing_ids = []
        next_job_id = str(max(existing_ids) + 1) if existing_ids else '1'
    else:
        next_job_id = '1'
    
    # Job data structure matching the web app format
    job_data = {
        'file_name': file_name,
        'file_size': file_size,
        'total_pages': total_pages,
        'status': 'pending',
        'created_at': now_ph,
        'local_path': file_path,
        'file_path': file_path, # Alias for GUI
        'upload_source': 'hotspot',
        'details': [
            {
                'file_name': file_name,
                'page_range': 'all',
                'color_mode': 'colored',
                'total_price': 0,
                'inserted_amount': 0,
                'change_amount': 0,
                'status': 'pending',
                'created_at': now_ph,
                'local_path': file_path,
                'file_path': file_path,
                'num_copies': 1,
                'page_size': 'Letter Size',
                'scale_mode': 'fit',
                'scale_percentage': 100
            }
        ]
    }
    
    # Save to Firebase with the auto-incremented job ID
    jobs_ref.child(next_job_id).set(job_data)
    job_data['job_id'] = next_job_id
    
    print(f"[Firebase] Job created: {next_job_id}, File: {file_name}, Pages: {total_pages}, Source: hotspot")
    return job_data


class SimpleHTTPRequestHandler(http.server.SimpleHTTPRequestHandler):
    """Custom HTTP request handler for file uploads and serving web pages."""
    
    upload_dir = None
    callback = None
    # Base path points to project root
    base_path = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    
    def do_GET(self):
        """Handle GET requests for serving HTML pages and static files."""
        if self.path == '/':
            # Serve main upload page
            self.send_response(200)
            self.send_header('Content-type', 'text/html')
            self.end_headers()
            
            html_path = os.path.join(self.base_path, 'templates', 'hotspot.html')
            try:
                with open(html_path, 'r', encoding='utf-8') as f:
                    html_content = f.read()
                self.wfile.write(html_content.encode())
            except FileNotFoundError:
                self.send_error(404, 'Template file not found: hotspot.html')
                
        elif self.path.startswith('/success'):
            # Serve success page
            self.send_response(200)
            self.send_header('Content-type', 'text/html')
            self.end_headers()
            
            success_path = os.path.join(self.base_path, 'templates', 'hotspot_uploaded.html')
            try:
                with open(success_path, 'r', encoding='utf-8') as f:
                    html_content = f.read()
                self.wfile.write(html_content.encode())
            except FileNotFoundError:
                self.send_error(404, 'Template file not found: hotspot_uploaded.html')
                
        elif self.path.startswith('/upload'):
            # Redirect GET requests for /upload to the main page
            self.send_response(302)
            self.send_header('Location', '/')
            self.end_headers()
            print("Redirecting GET request for /upload to /")
            
        elif self.path.startswith('/static/'):
            # Serve static files (CSS, JS, images)
            file_path = os.path.join(self.base_path, self.path[1:])
            try:
                with open(file_path, 'rb') as f:
                    content = f.read()
                
                # Determine content type
                if self.path.endswith('.css'):
                    content_type = 'text/css'
                elif self.path.endswith('.js'):
                    content_type = 'application/javascript'
                elif self.path.endswith('.jpg') or self.path.endswith('.jpeg'):
                    content_type = 'image/jpeg'
                elif self.path.endswith('.png'):
                    content_type = 'image/png'
                else:
                    content_type = 'application/octet-stream'
                
                self.send_response(200)
                self.send_header('Content-type', content_type)
                self.end_headers()
                self.wfile.write(content)
            except FileNotFoundError:
                self.send_response(404)
                self.end_headers()
                self.wfile.write(b'File not found')
        else:
            super().do_GET()
    
    def do_POST(self):
        """Handle POST requests for file uploads."""
        if self.path == '/upload':
            content_type = self.headers['Content-Type']
            if 'multipart/form-data' not in content_type:
                self.send_response(400)
                self.end_headers()
                return
            
            boundary = content_type.split('boundary=')[1].encode()
            remainbytes = int(self.headers['Content-Length'])
            line = self.rfile.readline()
            remainbytes -= len(line)
            
            if boundary not in line:
                self.send_response(400)
                self.end_headers()
                return
            
            line = self.rfile.readline()
            remainbytes -= len(line)
            
            filename = None
            for part in line.decode().split(';'):
                if 'filename=' in part:
                    filename = part.split('filename=')[1].strip().strip('"')
                    break
            
            if not filename:
                self.send_response(400)
                self.end_headers()
                return
            
            line = self.rfile.readline()
            remainbytes -= len(line)
            line = self.rfile.readline()
            remainbytes -= len(line)
            
            filepath = os.path.join(self.upload_dir, filename)
            base, ext = os.path.splitext(filepath)
            counter = 1
            while os.path.exists(filepath):
                filepath = f"{base}_{counter}{ext}"
                counter += 1
            
            try:
                with open(filepath, 'wb') as f:
                    preline = self.rfile.readline()
                    remainbytes -= len(preline)
                    
                    while remainbytes > 0:
                        line = self.rfile.readline()
                        remainbytes -= len(line)
                        
                        if boundary in line:
                            preline = preline[0:-1]
                            if preline.endswith(b'\r'):
                                preline = preline[0:-1]
                            f.write(preline)
                            break
                        else:
                            f.write(preline)
                            preline = line
                
                # Save to Firebase and get job data
                job_data = save_file_to_firebase(filepath, filename)
                
                # Call callback with both filepath and job data
                if self.callback:
                    self.callback(filepath, job_data)
                
                # Send a JSON response indicating success
                self.send_response(200)
                self.send_header('Content-type', 'application/json')
                self.end_headers()
                response = {'status': 'success', 'filename': filename}
                self.wfile.write(json.dumps(response).encode())
                print(f"Successfully uploaded {filename}")
                
            except Exception as e:
                print(f"Error during file upload: {e}")
                self.send_response(500)
                self.end_headers()


class HotspotServer:
    """Manages the HTTP server for hotspot file transfers."""
    
    def __init__(self, port=8080, upload_directory=None, callback=None):
        """
        Initialize the hotspot server.
        
        Args:
            port (int): Port number to run the server on
            upload_directory (str): Directory path where uploaded files will be saved
            callback (callable): Function to call when a file is received
        """
        self.port = port
        self.upload_directory = upload_directory or os.path.join(
            os.path.expanduser("~"), "Documents", "PrinTech_Received_Files"
        )
        self.callback = callback
        self.server = None
        self.server_thread = None
        
        # Create upload directory if it doesn't exist
        os.makedirs(self.upload_directory, exist_ok=True)
        
        # Configure the request handler
        SimpleHTTPRequestHandler.upload_dir = self.upload_directory
        SimpleHTTPRequestHandler.callback = self.callback
    
    def start(self):
        """Start the HTTP server in a separate thread."""
        try:
            self.server_thread = threading.Thread(target=self._run_server, daemon=True)
            self.server_thread.start()
            
            print(f"Server started automatically on port {self.port}")
            print(f"Files will be saved to: {self.upload_directory}")
            
        except Exception as e:
            print(f"Error starting server: {e}")
    
    def _run_server(self):
        """Internal method to run the HTTP server."""
        try:
            # Bind to all interfaces (0.0.0.0) so it can be accessed from any IP
            with socketserver.TCPServer(("0.0.0.0", self.port), SimpleHTTPRequestHandler) as httpd:
                self.server = httpd
                print(f"HTTP Server successfully bound to port {self.port}")
                print(f"Server accessible at:")
                print(f"  - http://localhost:{self.port}")
                print(f"  - http://192.168.137.1:{self.port}")
                print(f"  - Or any local network IP on port {self.port}")
                httpd.serve_forever()
        except Exception as e:
            print(f"Server error: {e}")
            print("This might be because another service is using port 8080")
    
    def stop(self):
        """Stop the HTTP server."""
        if self.server:
            print("Shutting down server...")
            self.server.shutdown()
            self.server = None
    
    def update_callback(self, callback):
        """
        Update the callback function for file uploads.
        
        Args:
            callback (callable): New callback function
        """
        self.callback = callback
        SimpleHTTPRequestHandler.callback = callback


class HotspotConfig:
    """Configuration settings for the hotspot."""
    
    # Default hotspot settings
    HOTSPOT_NAME = "PrinTech Hotspot"
    HOTSPOT_PASSWORD = "printech123"
    HOTSPOT_IP = "192.168.137.1"
    DEFAULT_PORT = 8080
    
    @classmethod
    def get_hotspot_url(cls, port=None):
        """Get the full hotspot URL."""
        port = port or cls.DEFAULT_PORT
        return f"http://{cls.HOTSPOT_IP}:{port}"
    
    @classmethod
    def get_upload_directory(cls):
        """Get the default upload directory path."""
        return os.path.join(os.path.expanduser("~"), "Documents", "PrinTech_Received_Files")
    
    @classmethod
    def generate_qr_code(cls, port=None, output_path=None):
        """
        Generate a QR code for the hotspot file upload URL.
        
        Args:
            port (int): Port number for the hotspot server
            output_path (str): Optional path to save the QR code image
            
        Returns:
            PIL.Image: QR code image
        """
        port = port or cls.DEFAULT_PORT
        url = f"http://{cls.HOTSPOT_IP}:{port}"
        
        # Create QR code
        qr = qrcode.QRCode(
            version=1,
            error_correction=qrcode.constants.ERROR_CORRECT_L,
            box_size=10,
            border=4,
        )
        qr.add_data(url)
        qr.make(fit=True)
        
        # Create image
        img = qr.make_image(fill_color="black", back_color="white")
        
        # Save to file if path provided
        if output_path:
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            img.save(output_path)
            print(f"QR code saved to: {output_path}")
        
        return img
